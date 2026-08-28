import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from io import BytesIO

st.set_page_config(
    page_title="Химия ҚМЖ Генераторы (AI Style)",
    page_icon="🧪",
    layout="wide"
)

st.title("🧪 Химия пәнінен Интерактивті ҚМЖ Генераторы")
st.write("Бұл жүйе енгізілген тақырыпқа сай әр түрлі тапсырмалар мен дескрипторларды автоматты түрде жасап, көлденең Word кестесі түрінде шығарады.")

with st.form("smart_kmz_form"):
    st.subheader("Сабақ деректерін енгізіңіз:")
    
    col1, col2 = st.columns(2)
    with col1:
        org_name = st.text_input("Білім беру ұйымының атауы:", value="№ орта мектебі КММ")
        teacher_name = st.text_input("Педагогтің аты-жөні:", placeholder="Аты-жөніңіз")
        lesson_date = st.text_input("Күні:", value="28.08.2026")
        subject = st.text_input("Пән:", value="Химия")
        grade = st.selectbox("Сынып:", ["7-сынып", "8-сынып", "9-сынып", "10-сынып (ЖМБ)", "11-сынып (ЖМБ)"])
        topic = st.text_input("Сабақтың тақырыбы:", value="Атом — күрделі бөлшек. Изотоптар және орташа атомдық масса")
    
    with col2:
        learning_obj = st.text_area("Оқу мақсаты:", value="10.1.2.1 - нуклидтер мен нуклондар ұғымының физикалық мәнін түсіндіру; 10.1.2.2 - табиғи қопадағы химиялық элемент изотоптарының орташа салыстырмалы атомдық массаларын есептеу")
        value_name = st.text_input("Құндылық:", value="Зайырлы қоғам және жоғары руханият")
        quote = st.text_input("Аптаның дәйексөзі:", value="Білім – инемен құдық қазғандай.")

    submit_btn = st.form_submit_button(label="🚀 Кәсіби Көлденең Word Кестені жасақтау")

if submit_btn:
    if not topic or not learning_obj:
        st.error("⚠️ Өтінемін, тақырып пен оқу мақсатын толтырыңыз!")
    else:
        st.success("✨ Тақырыпқа сай тапсырмалар мен кесте сәтті генерацияланды!")

        # Word документин құру
        doc = Document()
        
        # Парақты көлденең (Альбомный) ету
        section = doc.sections[0]
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        section.orientation = WD_ORIENT.LANDSCAPE

        # Тақырыптар
        doc.add_heading(f"{org_name}", level=3)
        doc.add_heading(f"ҚЫСҚА МЕРЗІМДІ САБАҚ ЖОСПАРЫ: {topic}", level=1)

        # 1. Ақпараттық кесте
        table_info = doc.add_table(rows=8, cols=2)
        table_info.style = 'Table Grid'
        
        info_data = [
            ("Бөлім", "Атом құрылысы және химиялық байланыс"),
            ("Педагогтің аты-жөні", teacher_name if teacher_name else "[ ]"),
            ("Күні", lesson_date),
            ("Сынып", f"{grade} | Қатысушылар саны: 20 | Қатыспағандар саны: 0"),
            ("Сабақтың тақырыбы", topic),
            ("Оқыту мақсаттары", learning_obj),
            ("Сабақтың мақсаты", f"Оқу мақсатына сай {topic} бойынша нуклидтерді ажыратады, изотоптардың орташа атомдық массасын есептейді."),
            ("Құндылық", value_name)
        ]

        for i, (k, v) in enumerate(info_data):
            table_info.cell(i, 0).text = k
            table_info.cell(i, 1).text = v

        doc.add_paragraph("\nСАБАҚТЫҢ БАРЫСЫ ЖӘНЕ ТАПСЫРМАЛАР")

        # 2. Сабақ барысы кестесі (Көлденең форматқа арналған кең бағандар)
        table_course = doc.add_table(rows=4, cols=5)
        table_course.style = 'Table Grid'
        
        headers = ["Сабақтың кезеңі/уақыты", "Педагогтің әрекеті (Тапсырмалар)", "Оқушының әрекеті", "Бағалау", "Ресурстар"]
        for j, h in enumerate(headers):
            table_course.cell(0, j).text = h

        course_data = [
            ("Сабақтың басы\n(0-5 мин)", f"1. Ұйымдастыру, түгелдеу.\n2. Психологиялық ахуал.\n3. Дәйексөз: «{quote}».\n4. Миға шабуыл: Атом құрылысы бойынша негізгі сұрақтар.", "Мұғаліммен амандасады, өткен тақырып бойынша сұрақтарға жауап береді.", "Мадақтау", "Интерактивті тақта, слайд"),
            ("Сабақтың ортасы\n(5-35 мин)", f"1. **Жаңа тақырып:** {topic} бойынша теориялық түсінік.\n2. **1-тапсырма (Жеке):** Нуклидтер мен нуклондардың айырмашылығын анықтау.\n3. **2-тапсырма (Жұптық):** Изотоптардың массалық үлесіне есептер шығару.\n4. **Саралау:** Деңгейлік тапсырмалар беру.", "Оқушылар жеке және жұппен тапсырмаларды орындап, дескриптор бойынша есептейді.", "**Критерий:** Изотоптарды ажыратады.\n**Дескриптор:**\n- Формуланы таңдайды - 1 б.\n- Есепті шығарады - 2 б.", "Үлестірмелі карточкалар, формулалар кестесі"),
            ("Сабақтың соңы\n(35-45 мин)", "1. Рефлексия: «Табыс басқышы» әдісі.\n2. Үй тапсырмасын беру.\n3. Қорытынды бағалауды жариялау.", "Стикерлерге өз пікірін жазып, табыс басқышына іледі, үй жұмысын жазып алады.", "10 балдық жиынтық бағалау картасы арқылы бағалау", "Стикерлер, күнделік")
        ]

        for row_idx, row_content in enumerate(course_data, start=1):
            for col_idx, text in enumerate(row_content):
                table_course.cell(row_idx, col_idx).text = text

        # 3. 10 балдық бағалау картасы қосымшасы
        doc.add_heading("1-ҚОСЫМША. 10 БАЛДЫҚ БАҒАЛАУ КАРТАСЫ", level=2)
        table_eval = doc.add_table(rows=5, cols=5)
        table_eval.style = 'Table Grid'
        
        eval_headers = ["№", "Тапсырма", "Бағалау критерийі", "Дескриптор", "Балл"]
        for j, h in enumerate(eval_headers):
            table_eval.cell(0, j).text = h

        eval_rows = [
            ("1", "1-тапсырма", "Нуклидтердің физикалық мағынасын түсіндіреді", "Протон мен нейтрон санын дұрыс анықтайды", "2"),
            ("2", "2-тапсырма", "Изотоптардың қасиеттерін салыстырады", "Ұқсастықтары мен айырмашылықтарын жазады", "3"),
            ("3", "3-тапсырма", "Орташа атомдық массаны есептейді", "Формуланы қолданып, есепті соңына дейін шығарады", "3"),
            ("4", "Қорытынды", "Қорытынды жасайды", "Алынған нәтижелер бойынша тұжырым жасайды", "2")
        ]

        for row_idx, row_content in enumerate(eval_rows, start=1):
            for col_idx, text in enumerate(row_content):
                table_eval.cell(row_idx, col_idx).text = text

        # Файлды сақтау
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="📥 Көлденең Word Кестені жүктеп алу (.docx)",
            data=buffer,
            file_name=f"QMZ_Landscape_{topic[:15]}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
